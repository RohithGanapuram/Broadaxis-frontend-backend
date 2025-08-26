#server
from mcp.server.fastmcp import FastMCP
import json
import os
import sys
from typing import List
import re
import logging
import traceback
from datetime import datetime
import tempfile

from pinecone import Pinecone
from tavily import TavilyClient
from dotenv import load_dotenv
# Load environment variables from parent directory
load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))
# File generation imports
from reportlab.lib.pagesizes import letter, A4, legal
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import uuid
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from io import BytesIO
from openai import OpenAI



# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('mcp_server.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('MCPServer')

# Error handling decorator
def handle_tool_errors(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            error_msg = f"Error in {func.__name__}: {str(e)}"
            logger.error(f"{error_msg}\n{traceback.format_exc()}")
            return json.dumps({
                "status": "error",
                "error": error_msg,
                "timestamp": datetime.now().isoformat()
            })
    return wrapper


client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
EMBED_MODEL = "text-embedding-3-small"  # 1536-dim
#Connection to Pinecone
try:
    pinecone_api_key = os.environ.get('PINECONE_API_KEY')
    if not pinecone_api_key:
        logger.warning("PINECONE_API_KEY not found in environment variables")
        pc = None
        index = None
    else:
        pc = Pinecone(api_key=pinecone_api_key)
        index = pc.Index("final-index")
        
        # Test the connection and check index stats
        try:
            stats = index.describe_index_stats()
            logger.info(f"Pinecone connection established. Index stats: {stats}")
            
            # Check if broadaxis-index namespace exists
            namespaces = stats.get('namespaces', {})
            if 'broadaxis-index' in namespaces:
                vector_count = namespaces['broadaxis-index'].get('vector_count', 0)
                logger.info(f"Found {vector_count} vectors in 'broadaxis-index' namespace")
            else:
                logger.warning("'broadaxis-index' namespace not found. Available namespaces: " + str(list(namespaces.keys())))
                
        except Exception as stats_error:
            logger.warning(f"Could not get index stats: {stats_error}")
            
except Exception as e:
    logger.error(f"Failed to initialize Pinecone: {e}")
    pc = None
    index = None

def _embed_text(text: str):
    """Return a single OpenAI embedding vector for the given text."""
    try:
        # Ensure text is not empty and clean it
        if not text or not text.strip():
            raise ValueError("Text cannot be empty")
        
        clean_text = text.strip()
        logger.info(f"Embedding text of length: {len(clean_text)}")
        
        # Enhance query for better semantic matching
        enhanced_query = clean_text
        if any(word in clean_text.lower() for word in ["job", "duties", "responsibilities", "role"]):
            enhanced_query = f"Job responsibilities and duties: {clean_text}"
        elif any(word in clean_text.lower() for word in ["service", "capability", "expertise", "what does"]):
            enhanced_query = f"Company services and capabilities: {clean_text}"
        
        # Create high-quality embedding
        resp = client.embeddings.create(
            model=EMBED_MODEL, 
            input=[enhanced_query]
            # dimensions parameter not needed - uses model default (1536)
        )
        embedding = resp.data[0].embedding
        
        logger.info(f"Generated embedding with dimension: {len(embedding)}")
        return embedding  # list[float], length 1536
        
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        raise

def sanitize_filename(name: str) -> str:
    # Remove all non-alphanumeric, dash, underscore characters
    return re.sub(r'[^a-zA-Z0-9-_]', '_', name)

def process_inline_formatting(text: str) -> str:
    """
    Process inline markdown formatting (bold, italic) in text.
    Note: This is a simplified version - for full formatting, consider using python-docx's Run objects for rich formatting.
    """
    # For now, we'll just return the text as-is
    # In a full implementation, you could use python-docx's Run objects for rich formatting
    return text

def get_file_type_info(file_path: str) -> dict:
    """
    Get file type information and validation details.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file type information
    """
    file_extension = file_path.lower().split('.')[-1] if '.' in file_path else ''
    
    # Define file type categories
    text_extensions = ['txt', 'md', 'json', 'xml', 'csv', 'log', 'py', 'js', 'html', 'css', 'yaml', 'yml', 'ini', 'cfg']
    document_extensions = ['pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx']
    image_extensions = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'svg', 'webp']
    archive_extensions = ['zip', 'rar', '7z', 'tar', 'gz']
    
    is_text_file = file_extension in text_extensions
    is_document = file_extension in document_extensions
    is_image = file_extension in image_extensions
    is_archive = file_extension in archive_extensions
    
    return {
        'extension': file_extension,
        'is_text_file': is_text_file,
        'is_document': is_document,
        'is_image': is_image,
        'is_archive': is_archive,
        'mime_type': _get_mime_type(file_extension),
        'recommended_max_size_mb': _get_recommended_max_size(file_extension)
    }

def _get_mime_type(extension: str) -> str:
    """Get MIME type for file extension."""
    mime_types = {
        'txt': 'text/plain',
        'md': 'text/markdown',
        'json': 'application/json',
        'xml': 'application/xml',
        'csv': 'text/csv',
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xls': 'application/vnd.ms-excel',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png',
        'gif': 'image/gif',
        'zip': 'application/zip',
        'rar': 'application/vnd.rar'
    }
    return mime_types.get(extension, 'application/octet-stream')

def _get_recommended_max_size(extension: str) -> int:
    """Get recommended maximum file size in MB for file type."""
    size_limits = {
        'txt': 10,
        'md': 10,
        'json': 50,
        'xml': 50,
        'csv': 100,
        'pdf': 100,
        'doc': 50,
        'docx': 50,
        'xls': 50,
        'xlsx': 50,
        'jpg': 50,
        'jpeg': 50,
        'png': 50,
        'gif': 50,
        'zip': 200,
        'rar': 200
    }
    return size_limits.get(extension, 50)  # Default 50MB

def _calculate_relevance_score(file_info: dict, query: str, search_type: str) -> float:
    """
    Calculate relevance score for search results.
    
    Args:
        file_info: File information dictionary
        query: Search query
        search_type: Type of search performed
        
    Returns:
        Relevance score between 0.0 and 1.0
    """
    score = 0.0
    query_lower = query.lower()
    file_name = file_info.get('name', '').lower()
    
    if search_type.lower() == "filename":
        # Exact match gets highest score
        if query_lower == file_name:
            score = 1.0
        elif query_lower in file_name:
            score = 0.8
        elif any(word in file_name for word in query_lower.split()):
            score = 0.6
        else:
            score = 0.2  # Partial match
    
    elif search_type.lower() == "content":
        # Content search relevance (would need content analysis)
        matches = file_info.get('search_matches', [])
        if matches:
            score = min(len(matches) * 0.2, 1.0)  # More matches = higher score
        else:
            score = 0.1
    
    else:  # "both"
        # Combine filename and content scores
        filename_score = _calculate_relevance_score(file_info, query, "filename")
        content_score = _calculate_relevance_score(file_info, query, "content")
        score = (filename_score * 0.6) + (content_score * 0.4)  # Weight filename more
    
    return round(score, 3)

def _get_content_preview(file_path: str, query: str, max_length: int = 200) -> str:
    """
    Get content preview showing search matches.
    
    Args:
        file_path: Path to the file
        query: Search query
        max_length: Maximum preview length
        
    Returns:
        Content preview string
    """
    try:
        # This would need to be implemented based on your SharePoint manager capabilities
        # For now, return a placeholder
        return f"Content preview for '{query}' in {file_path} (preview not implemented)"
    except Exception as e:
        logger.warning(f"Could not get content preview: {e}")
        return ""

def _merge_search_results(filename_result: dict, content_result: dict, query: str) -> dict:
    """
    Merge filename and content search results.
    
    Args:
        filename_result: Results from filename search
        content_result: Results from content search
        query: Original search query
        
    Returns:
        Merged results dictionary
    """
    # This is a placeholder - would need to be implemented based on your SharePoint manager
    # For now, prioritize filename results and add content results
    merged_files = []
    
    if filename_result.get('status') == 'success':
        merged_files.extend(filename_result.get('files', []))
    
    if content_result.get('status') == 'success':
        # Add content results that aren't already in filename results
        filename_paths = {f.get('path', '') for f in merged_files}
        for content_file in content_result.get('files', []):
            if content_file.get('path', '') not in filename_paths:
                merged_files.append(content_file)
    
    return {
        'status': 'success',
        'files': merged_files
    }

def _parse_page_range(page_spec: str, total_pages: int, max_pages: int) -> list:
    """
    Parse page range specification and return list of page numbers to extract.
    
    Args:
        page_spec: Page specification string
        total_pages: Total number of pages in PDF
        max_pages: Maximum pages allowed to extract
        
    Returns:
        List of page numbers to extract
    """
    if not page_spec or page_spec.lower() == "all":
        # Extract all pages up to max_pages
        return list(range(1, min(total_pages + 1, max_pages + 1)))
    
    if page_spec.lower() == "first":
        return [1]
    
    if page_spec.lower() == "last":
        return [total_pages]
    
    pages = []
    
    # Handle comma-separated pages
    if ',' in page_spec:
        for part in page_spec.split(','):
            part = part.strip()
            if '-' in part:
                # Handle ranges like "1-5"
                start, end = part.split('-', 1)
                try:
                    start_page = int(start.strip())
                    end_page = int(end.strip())
                    pages.extend(range(start_page, end_page + 1))
                except ValueError:
                    continue
            else:
                # Handle single pages
                try:
                    page_num = int(part)
                    pages.append(page_num)
                except ValueError:
                    continue
    
    # Handle single range like "1-5"
    elif '-' in page_spec:
        start, end = page_spec.split('-', 1)
        try:
            start_page = int(start.strip())
            end_page = int(end.strip())
            pages = list(range(start_page, end_page + 1))
        except ValueError:
            return []
    
    # Handle single page number
    else:
        try:
            page_num = int(page_spec.strip())
            pages = [page_num]
        except ValueError:
            return []
    
    # Filter valid pages and apply max_pages limit
    valid_pages = [p for p in pages if 1 <= p <= total_pages]
    return valid_pages[:max_pages]

def _clean_pdf_text(text: str) -> str:
    """
    Clean and format extracted PDF text.
    
    Args:
        text: Raw extracted text from PDF
        
    Returns:
        Cleaned and formatted text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove common PDF artifacts
    text = re.sub(r'[^\w\s\-.,;:!?()\[\]{}"\']+', '', text)
    
    # Fix common OCR issues
    text = text.replace('|', 'I')  # Common OCR mistake
    text = text.replace('0', 'O')  # In certain contexts
    
    # Remove page numbers and headers/footers
    text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)  # Standalone page numbers
    
    # Clean up line breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple line breaks to double line breaks
    
    return text.strip()

def _preserve_document_structure(text: str) -> str:
    """
    Preserve document structure like headings and sections.
    
    Args:
        text: Cleaned PDF text
        
    Returns:
        Text with preserved structure
    """
    if not text:
        return ""
    
    lines = text.split('\n')
    structured_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect headings (all caps, short lines)
        if line.isupper() and len(line) < 100:
            structured_lines.append(f"\n## {line.title()}\n")
        # Detect section headers (numbered items)
        elif re.match(r'^\d+\.\s+[A-Z]', line):
            structured_lines.append(f"\n### {line}\n")
        # Detect bullet points
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            structured_lines.append(f"  {line}")
        else:
            structured_lines.append(line)
    
    return '\n'.join(structured_lines)

def _extract_table_data(page) -> list:
    """
    Extract table data from PDF page.
    
    Args:
        page: PDF page object
        
    Returns:
        List of table data (placeholder implementation)
    """
    # This is a placeholder - would need more sophisticated table extraction
    # For now, return empty list
    # In a full implementation, you could use libraries like tabula-py or camelot-py
    return []

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


#create an MCP server
mcp = FastMCP("Research-Demo")

#Add an addition tool
@mcp.tool()
def sum(a: int, b: int) -> int:
    """Add two numbers"""
    return a+b


@mcp.tool()
def Broadaxis_knowledge_search(query: str, top_k: int = 5, min_score: float = 0.2, include_scores: bool = False, source_filter: str = None):
    """
    Retrieves the most relevant company's (Broadaxis) information from the internal knowledge base.
    Semantic search over Pinecone using OpenAI embeddings.
    
    Args:
        query: Search query string
        top_k: Number of top results to return (default: 5)
        min_score: Minimum similarity score threshold (default: 0.7)
        include_scores: Whether to include similarity scores in response (default: False)
        source_filter: Optional source filter (e.g., "proposals", "case_studies") (default: None)
    """
    if not query or not query.strip():
        return json.dumps({"error": "Query cannot be empty"})

    # Make sure we have an index
    if not index:
        return json.dumps({"error": "Pinecone index not available"})

    try:
        logger.info(f"Searching for: {query.strip()}")
        
        # Step 1: Embed the query with OpenAI
        query_embedding = _embed_text(query.strip())
        logger.info(f"Query embedding dimension: {len(query_embedding)}")

        # Step 2: Query Pinecone with optimized parameters
        query_result = index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True,
            namespace="broadaxis-index",
            filter={}  # Can add filters if needed
        )
        
        logger.info(f"Pinecone query result type: {type(query_result)}")
        logger.info(f"Raw query result: {query_result}")

        # Step 3: Extract matches properly
        matches = []
        if hasattr(query_result, 'matches'):
            matches = query_result.matches
        elif isinstance(query_result, dict) and 'matches' in query_result:
            matches = query_result['matches']
        
        logger.info(f"Found {len(matches)} matches")

        # Step 4: Extract and return clean text content for LLM
        documents = []
        scored_documents = []
        
        for i, match in enumerate(matches):
            try:
                # Get metadata
                metadata = {}
                if hasattr(match, 'metadata'):
                    metadata = match.metadata or {}
                elif isinstance(match, dict) and 'metadata' in match:
                    metadata = match['metadata'] or {}
                
                # Get score for filtering and logging
                score = 0.0
                if hasattr(match, 'score'):
                    score = match.score
                elif isinstance(match, dict) and 'score' in match:
                    score = match['score']
                
                # Get text content
                text_content = metadata.get('text', '')
                source = metadata.get('source', 'Unknown')
                
                # Apply source filter if specified
                if source_filter and source_filter.lower() not in source.lower():
                    logger.info(f"Match {i+1} filtered out: source '{source}' doesn't match filter '{source_filter}'")
                    continue
                
                if text_content and score >= min_score:
                    # Clean text for LLM consumption
                    clean_text = text_content.replace('\uf0b7', '•').replace('\u2022', '•')
                    documents.append(clean_text)
                    
                    if include_scores:
                        scored_documents.append({
                            'text': clean_text,
                            'score': score,
                            'source': source
                        })
                    
                    logger.info(f"Match {i+1}: score={score:.4f}, source={source}, text_len={len(text_content)}")
                elif text_content:
                    logger.info(f"Match {i+1} filtered out: score={score:.4f} < {min_score}, source={source}")
                
            except Exception as match_error:
                logger.error(f"Error processing match {i}: {match_error}")
                continue

        if documents:
            if include_scores:
                # Return structured response with scores
                return json.dumps({
                    "status": "success",
                    "query": query.strip(),
                    "results": scored_documents,
                    "total_results": len(scored_documents),
                    "min_score_used": min_score
                })
            else:
                # Return clean text content separated by double newlines (original format)
                return "\n\n".join(documents)
        else:
            return "No relevant information found in the BroadAxis knowledge base."

    except Exception as e:
        logger.error(f"Knowledge search error: {e}")
        logger.error(f"Error traceback: {traceback.format_exc()}")
        return f"Error searching BroadAxis knowledge base: {str(e)}"

# Initialize Tavily client with proper error handling
try:
    tavily_api_key = os.environ.get('TAVILY_API_KEY')
    if not tavily_api_key:
        logger.warning("TAVILY_API_KEY not found in environment variables")
        tavily = None
    else:
        tavily = TavilyClient(api_key=tavily_api_key)
        logger.info("Tavily client initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Tavily client: {e}")
    tavily = None

@mcp.tool()
def web_search_tool(query: str):
    """
    Performs a real-time web search using Tavily and returns relevant results
    (including title, URL, and snippet).

    Args:
        query: A natural language search query.

    Returns:
        A JSON string with the top search results.
    """
    try:
        if not query or not query.strip():
            return json.dumps({"error": "Search query cannot be empty"})
        
        # Check if Tavily client is available
        if tavily is None:
            return json.dumps({
                "error": "Web search is currently unavailable. Tavily API key not configured or invalid.",
                "status": "unavailable"
            })
        
        logger.info(f"Performing web search for query: {query.strip()}")
        results = tavily.search(query=query.strip(), search_depth="advanced", include_answer=False)
        
        if not results or "results" not in results:
            return json.dumps({
                "error": "No search results found",
                "status": "no_results"
            })
        
        formatted = [
            {
                "title": r.get("title", "No title"),
                "url": r.get("url", ""),
                "snippet": r.get("content", "No content")
            }
            for r in results.get("results", [])
        ]

        logger.info(f"Web search completed successfully. Found {len(formatted)} results")
        return json.dumps({
            "results": formatted, 
            "total_results": len(formatted),
            "status": "success"
        })

    except Exception as e:
        logger.error(f"Web search error: {e}")
        return json.dumps({
            "error": f"Search failed: {str(e)}",
            "status": "error"
        })
    
@mcp.tool()
def generate_pdf_document(title: str, content: str, filename: str = None, page_size: str = "letter", include_toc: bool = False) -> str:
    """
    Generate and create a professional PDF document from text content with markdown support.
    Automatically uploads to SharePoint for storage and sharing.
    
    Use this tool when the user asks to create, generate, or make a PDF file.

    Args:
        title: The title of the document (appears as centered heading)
        content: The main content of the document with markdown formatting support:
            - "# " → Heading 1
            - "## " → Heading 2  
            - "### " → Heading 3
            - "- " or "* " → Bullet points
            - "1. " → Numbered lists
            - Regular text → Normal paragraphs
        filename: Optional custom filename (without extension, will be sanitized)
        page_size: Page size format ("letter", "a4", "legal") - default: "letter"
        include_toc: Whether to include a table of contents for documents with multiple headings

    Returns:
        JSON string with file generation status including:
        - Success: SharePoint path, file size, creation timestamp
        - Fallback: Local file path if SharePoint unavailable
        - Error: Detailed error message with cleanup performed

    Examples:
        Basic usage:
        generate_pdf_document("Project Proposal", "# Overview\n\nThis is a proposal...")
        
        With custom filename:
        generate_pdf_document("Technical Spec", "## Requirements\n\n- Cloud infrastructure", "tech_spec_2024")
        
        With table of contents:
        generate_pdf_document("User Manual", "# Chapter 1\n## Section 1.1\n...", include_toc=True)
    """
    logger.info(f"Starting PDF generation: {title}")
    temp_file_path = None
    
    # Input validation
    if not title or not title.strip():
        return json.dumps({"status": "error", "error": "Title cannot be empty"})
    
    if not content or not content.strip():
        return json.dumps({"status": "error", "error": "Content cannot be empty"})
    
    # Validate page size
    valid_page_sizes = ["letter", "a4", "legal"]
    if page_size.lower() not in valid_page_sizes:
        return json.dumps({"status": "error", "error": f"Invalid page size. Must be one of: {', '.join(valid_page_sizes)}"})
    
    try:
        # Generate unique filename if not provided
        if not filename:
            filename = f"document_{uuid.uuid4().hex[:8]}"

        # Ensure filename doesn't have extension
        filename = sanitize_filename(filename.replace('.pdf', ''))

        # Create output directory if it doesn't exist
        output_dir = os.path.join(BASE_DIR, "generated_files")
        os.makedirs(output_dir, exist_ok=True)
        
        # Create file path
        temp_file_path = os.path.join(output_dir, f"{filename}.pdf")
        
        # Create PDF document with configurable page size
        page_sizes = {
            "letter": letter,
            "a4": A4,
            "legal": legal
        }
        selected_page_size = page_sizes.get(page_size.lower(), letter)
        
        doc = SimpleDocTemplate(temp_file_path, pagesize=selected_page_size)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Initialize headings_for_toc for table of contents
        headings_for_toc = []  # Track headings for table of contents
        
        # First pass: collect headings for table of contents
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if line:
                if line.startswith('# '):
                    heading_text = line[2:]
                    if include_toc:
                        headings_for_toc.append((1, heading_text))
                elif line.startswith('## '):
                    heading_text = line[3:]
                    if include_toc:
                        headings_for_toc.append((2, heading_text))
                elif line.startswith('### '):
                    heading_text = line[4:]
                    if include_toc:
                        headings_for_toc.append((3, heading_text))

        # Add table of contents if requested and headings exist
        if include_toc and headings_for_toc:
            toc_style = ParagraphStyle(
                'TOC',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=20,
                alignment=0
            )
            story.append(Paragraph("Table of Contents", toc_style))
            story.append(Spacer(1, 10))
            
            for level, heading in headings_for_toc:
                indent = "&nbsp;" * (level - 1) * 4  # Indent based on heading level
                story.append(Paragraph(f"{indent}• {heading}", styles['Normal']))
                story.append(Spacer(1, 3))
            
            story.append(Spacer(1, 20))

        # Second pass: process content with enhanced markdown support
        for line in content_lines:
            line = line.strip()
            if line:
                # Escape special characters for ReportLab
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                if line.startswith('# '):
                    heading_text = line[2:]
                    story.append(Paragraph(heading_text, styles['Heading1']))
                elif line.startswith('## '):
                    heading_text = line[3:]
                    story.append(Paragraph(heading_text, styles['Heading2']))
                elif line.startswith('### '):
                    heading_text = line[4:]
                    story.append(Paragraph(heading_text, styles['Heading3']))
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Handle numbered lists (basic support)
                    story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))

        # Build PDF
        doc.build(story)
        
        # Get file size
        file_size = os.path.getsize(temp_file_path)
        
        # Upload to SharePoint
        try:
            sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
            from sharepoint_api import SharePointManager  # type: ignore
            
            with open(temp_file_path, 'rb') as f:
                file_content = f.read()
            
            sharepoint_manager = SharePointManager()
            sharepoint_folder = f"Generated_Documents/{datetime.datetime.now().strftime('%Y-%m')}"
            
            logger.info(f"Attempting SharePoint upload: {filename}.pdf to {sharepoint_folder}")
            upload_result = sharepoint_manager.upload_file_to_sharepoint(
                file_content, f"{filename}.pdf", sharepoint_folder
            )
            logger.info(f"SharePoint upload result: {upload_result}")
            
            if upload_result['status'] == 'success':
                os.unlink(temp_file_path)  # Only delete if SharePoint upload succeeds
                return json.dumps({
                    "status": "success",
                    "filename": f"{filename}.pdf",
                    "sharepoint_path": sharepoint_folder,
                    "file_size": file_size,
                    "created_at": datetime.datetime.now().isoformat(),
                    "type": "pdf",
                    "page_size": page_size,
                    "has_toc": include_toc and len(headings_for_toc) > 0,
                    "heading_count": len(headings_for_toc),
                    "message": "PDF generated and saved to SharePoint"
                })
            else:
                logger.warning(f"SharePoint upload failed: {upload_result.get('message', 'Unknown error')}")
        except Exception as e:
            logger.warning(f"SharePoint upload failed: {e}")
        
        # Return success with local file path (SharePoint failed)
        return json.dumps({
            "status": "success",
            "filename": f"{filename}.pdf",
            "file_path": temp_file_path,
            "file_size": file_size,
            "created_at": datetime.datetime.now().isoformat(),
            "type": "pdf",
            "page_size": page_size,
            "has_toc": include_toc and len(headings_for_toc) > 0,
            "heading_count": len(headings_for_toc),
            "message": "PDF generated locally (SharePoint unavailable)"
        })

    except Exception as e:
        logger.error(f"PDF generation failed: {str(e)}")
        # Clean up on error
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except:
                pass
        return json.dumps({
            "status": "error",
            "error": str(e),
            "message": "Failed to generate PDF document"
        })


@mcp.tool()
def generate_word_document(title: str, content: str, filename: str = None, include_toc: bool = False, page_orientation: str = "portrait") -> str:
    """
    Generate and create a professional Word document from text content with enhanced markdown support.
    Automatically uploads to SharePoint for storage and sharing.
    
    Use this tool when the user asks to create, generate, or make a Word document.

    Args:
        title: The title of the document (appears as centered heading)
        content: The main content of the document with markdown formatting support:
            - "# " → Heading 1
            - "## " → Heading 2  
            - "### " → Heading 3
            - "#### " → Heading 4
            - "- " or "* " → Bullet points
            - "1. ", "2. ", "3. " → Numbered lists
            - "**text**" → Bold text
            - "*text*" → Italic text
            - Regular text → Normal paragraphs
        filename: Optional custom filename (without extension, will be sanitized)
        include_toc: Whether to include a table of contents for documents with multiple headings
        page_orientation: Page orientation ("portrait" or "landscape") - default: "portrait"

    Returns:
        JSON string with file generation status including:
        - Success: SharePoint path, file size, creation timestamp
        - Fallback: Local file path if SharePoint unavailable
        - Error: Detailed error message with cleanup performed

    Examples:
        Basic usage:
        generate_word_document("Project Proposal", "# Overview\n\nThis is a proposal...")
        
        With custom filename:
        generate_word_document("Technical Spec", "## Requirements\n\n- Cloud infrastructure", "tech_spec_2024")
        
        With table of contents:
        generate_word_document("User Manual", "# Chapter 1\n## Section 1.1\n...", include_toc=True)
        
        With landscape orientation:
        generate_word_document("Wide Table Report", "content...", page_orientation="landscape")
    """
    # Input validation
    if not title or not title.strip():
        return json.dumps({"status": "error", "error": "Title cannot be empty"})
    
    if not content or not content.strip():
        return json.dumps({"status": "error", "error": "Content cannot be empty"})
    
    # Validate page orientation
    valid_orientations = ["portrait", "landscape"]
    if page_orientation.lower() not in valid_orientations:
        return json.dumps({"status": "error", "error": f"Invalid page orientation. Must be one of: {', '.join(valid_orientations)}"})
    
    try:
        # Generate unique filename if not provided
        if not filename:
            filename = f"document_{uuid.uuid4().hex[:8]}"

        # Ensure filename doesn't have extension and sanitize it
        filename = sanitize_filename(filename.replace('.docx', '').replace('.doc', ''))

        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        
        # Create Word document with page orientation
        doc = Document()
        
        # Set page orientation if landscape
        if page_orientation.lower() == "landscape":
            section = doc.sections[0]
            section.orientation = 1  # 1 = landscape, 0 = portrait

        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some space after title
        doc.add_paragraph()
        
        # Initialize headings_for_toc for table of contents
        headings_for_toc = []  # Track headings for table of contents
        
        # First pass: collect headings for table of contents
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if line:
                if line.startswith('# '):
                    heading_text = line[2:]
                    if include_toc:
                        headings_for_toc.append((1, heading_text))
                elif line.startswith('## '):
                    heading_text = line[3:]
                    if include_toc:
                        headings_for_toc.append((2, heading_text))
                elif line.startswith('### '):
                    heading_text = line[4:]
                    if include_toc:
                        headings_for_toc.append((3, heading_text))
                elif line.startswith('#### '):
                    heading_text = line[5:]
                    if include_toc:
                        headings_for_toc.append((4, heading_text))

        # Add table of contents if requested and headings exist
        if include_toc and headings_for_toc:
            toc_heading = doc.add_heading("Table of Contents", level=2)
            doc.add_paragraph()
            
            for level, heading in headings_for_toc:
                indent = "  " * (level - 1)  # Indent based on heading level
                toc_text = f"{indent}• {heading}"
                doc.add_paragraph(toc_text, style='Normal')
            
            doc.add_paragraph()  # Add space after TOC

        # Second pass: process content with enhanced markdown support
        for line in content_lines:
            line = line.strip()
            if line:
                if line.startswith('# '):
                    heading_text = line[2:]
                    doc.add_heading(heading_text, level=1)
                elif line.startswith('## '):
                    heading_text = line[3:]
                    doc.add_heading(heading_text, level=2)
                elif line.startswith('### '):
                    heading_text = line[4:]
                    doc.add_heading(heading_text, level=3)
                elif line.startswith('#### '):
                    heading_text = line[5:]
                    doc.add_heading(heading_text, level=4)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Handle numbered lists (basic support)
                    doc.add_paragraph(line, style='List Number')
                else:
                    # Process inline formatting (bold, italic)
                    processed_line = process_inline_formatting(line)
                    doc.add_paragraph(processed_line)
            else:
                doc.add_paragraph()

        # Save to temp file
        doc.save(temp_file.name)
        
        # Read file content for SharePoint upload
        with open(temp_file.name, 'rb') as f:
            file_content = f.read()
        
        os.unlink(temp_file.name)
        
        # Upload to SharePoint
        try:
            sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
            from sharepoint_api import SharePointManager  # type: ignore
            
            sharepoint_manager = SharePointManager()
            sharepoint_folder = f"Generated_Documents/{datetime.datetime.now().strftime('%Y-%m')}"
            
            upload_result = sharepoint_manager.upload_file_to_sharepoint(
                file_content, f"{filename}.docx", sharepoint_folder
            )
            
            if upload_result['status'] == 'success':
                return json.dumps({
                    "status": "success",
                    "filename": f"{filename}.docx",
                    "sharepoint_path": sharepoint_folder,
                    "file_size": len(file_content),
                    "created_at": datetime.datetime.now().isoformat(),
                    "type": "docx",
                    "page_orientation": page_orientation,
                    "has_toc": include_toc and len(headings_for_toc) > 0,
                    "heading_count": len(headings_for_toc),
                    "message": "Word document generated and saved to SharePoint"
                })
        except Exception as e:
            logger.warning(f"SharePoint upload failed: {e}")
        
        return json.dumps({
            "status": "success",
            "filename": f"{filename}.docx",
            "file_size": len(file_content),
            "created_at": datetime.datetime.now().isoformat(),
            "type": "docx",
            "page_orientation": page_orientation,
            "has_toc": include_toc and len(headings_for_toc) > 0,
            "heading_count": len(headings_for_toc),
            "message": "Word document generated locally (SharePoint unavailable)"
        })

    except Exception as e:
        logger.error(f"Word document generation failed: {str(e)}")
        return json.dumps({
            "status": "error",
            "error": str(e),
            "message": "Failed to generate Word document"
        })


@mcp.tool()
def sharepoint_read_file(path: str, max_size_mb: int = 50, encoding: str = "utf-8", preview_lines: int = 0) -> str:
    """
    Read a file from SharePoint with enhanced features and validation.
    
    Args:
        path: SharePoint file path (e.g., "Documents/file.txt", "RFP/proposal.pdf")
        max_size_mb: Maximum file size in MB to prevent memory issues (default: 50MB)
        encoding: Text encoding for text files (default: "utf-8")
        preview_lines: Number of lines to return for preview (0 = full file, default: 0)
    
    Returns:
        JSON string with file content and metadata or detailed error information
        
    Examples:
        Basic usage:
        sharepoint_read_file("Documents/proposal.txt")
        
        With preview:
        sharepoint_read_file("Documents/report.txt", preview_lines=10)
        
        Large file with custom size limit:
        sharepoint_read_file("Documents/manual.pdf", max_size_mb=100)
        
        Custom encoding:
        sharepoint_read_file("Documents/legacy.txt", encoding="latin-1")
    """
    # Input validation
    if not path or not path.strip():
        return json.dumps({
            "status": "error",
            "error": "File path cannot be empty",
            "error_code": "INVALID_PATH",
            "suggestions": ["Provide a valid file path"]
        })
    
    # Sanitize and validate path
    clean_path = path.strip().replace('\\', '/')
    if '..' in clean_path or clean_path.startswith('/'):
        return json.dumps({
            "status": "error", 
            "error": "Invalid file path - path traversal not allowed",
            "error_code": "PATH_TRAVERSAL",
            "path": clean_path,
            "suggestions": ["Use relative paths without '..' or absolute paths"]
        })
    
    # Validate parameters
    if max_size_mb <= 0 or max_size_mb > 500:
        return json.dumps({
            "status": "error",
            "error": "Invalid max_size_mb - must be between 1 and 500",
            "error_code": "INVALID_SIZE_LIMIT",
            "max_size_mb": max_size_mb
        })
    
    if preview_lines < 0:
        return json.dumps({
            "status": "error",
            "error": "preview_lines must be non-negative",
            "error_code": "INVALID_PREVIEW_LINES",
            "preview_lines": preview_lines
        })
    
    try:
        sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
        from sharepoint_api import SharePointManager  # type: ignore
        
        sharepoint_manager = SharePointManager()
        
        # Get file metadata first
        logger.info(f"Reading file from SharePoint: {clean_path}")
        start_time = datetime.datetime.now()
        
        result = sharepoint_manager.get_file_content(clean_path)
        
        if result['status'] == 'success':
            content = result['content']
            content_size = len(content)
            
            # Check file size limit
            if content_size > max_size_mb * 1024 * 1024:
                return json.dumps({
                    "status": "error",
                    "error": f"File size ({content_size / (1024*1024):.1f}MB) exceeds limit ({max_size_mb}MB)",
                    "error_code": "FILE_TOO_LARGE",
                    "path": clean_path,
                    "file_size_mb": content_size / (1024*1024),
                    "max_size_mb": max_size_mb,
                    "suggestions": [
                        f"Increase max_size_mb parameter to {max_size_mb * 2} or higher",
                        "Use preview_lines parameter to read only first few lines"
                    ]
                })
            
            # Get file type information
            file_info = get_file_type_info(clean_path)
            file_extension = file_info['extension']
            is_text_file = file_info['is_text_file']
            
            # Process content based on file type and preview settings
            processed_content = content
            if is_text_file and isinstance(content, bytes):
                try:
                    processed_content = content.decode(encoding)
                except UnicodeDecodeError:
                    # Try UTF-8 as fallback
                    try:
                        processed_content = content.decode('utf-8')
                    except UnicodeDecodeError:
                        processed_content = content.decode('utf-8', errors='replace')
            
            # Handle preview mode for text files
            if preview_lines > 0 and is_text_file:
                lines = processed_content.split('\n')
                if len(lines) > preview_lines:
                    processed_content = '\n'.join(lines[:preview_lines])
                    preview_note = f" (showing first {preview_lines} lines of {len(lines)} total)"
                else:
                    preview_note = ""
            # Handle preview mode for PDF files
            elif preview_lines > 0 and file_extension.lower() == 'pdf':
                try:
                    # Use extract_pdf_text to get a preview of the PDF
                    # Estimate how many pages to extract based on requested lines (roughly 50 lines per page)
                    estimated_pages = max(1, min(3, (preview_lines + 49) // 50))  # At least 1 page, max 3 pages
                    
                    pdf_preview_result = extract_pdf_text(clean_path, pages="first", max_pages=estimated_pages)
                    
                    # Parse the JSON result
                    pdf_data = json.loads(pdf_preview_result)
                    if pdf_data['status'] == 'success':
                        extracted_text = pdf_data.get('text', '')  # Use 'text' field from extract_pdf_text response
                        total_pages = pdf_data.get('total_pages', 0)
                        pages_extracted = pdf_data.get('pages_extracted', 0)
                        
                        # Limit to requested number of lines
                        lines = extracted_text.split('\n')
                        if len(lines) > preview_lines:
                            processed_content = '\n'.join(lines[:preview_lines])
                            preview_note = f" (showing first {preview_lines} lines from first {pages_extracted} pages of {total_pages} total pages)"
                        else:
                            processed_content = extracted_text
                            preview_note = f" (showing first {pages_extracted} pages of {total_pages} total pages)"
                    else:
                        # If PDF extraction fails, return error message
                        processed_content = f"Error extracting PDF preview: {pdf_data.get('error', 'Unknown error')}"
                        preview_note = " (PDF preview extraction failed)"
                        
                except Exception as pdf_error:
                    logger.warning(f"PDF preview extraction failed: {pdf_error}")
                    processed_content = f"Error extracting PDF preview: {str(pdf_error)}"
                    preview_note = " (PDF preview extraction failed)"
            else:
                preview_note = ""
            
            # Calculate processing time
            end_time = datetime.datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            # Build response
            response = {
                "status": "success",
                "content": processed_content,
                "path": clean_path,
                "size": content_size,
                "size_mb": round(content_size / (1024*1024), 2),
                "file_type": file_extension,
                "mime_type": file_info['mime_type'],
                "is_text_file": is_text_file,
                "is_document": file_info['is_document'],
                "is_image": file_info['is_image'],
                "is_archive": file_info['is_archive'],
                "encoding": encoding if is_text_file else None,
                "processing_time_seconds": round(processing_time, 3),
                "preview_mode": preview_lines > 0,
                "preview_lines": preview_lines if preview_lines > 0 else None,
                "recommended_max_size_mb": file_info['recommended_max_size_mb'],
                "timestamp": datetime.datetime.now().isoformat()
            }
            
            if preview_note:
                response["preview_note"] = preview_note
            
            return json.dumps(response)
        else:
            # Enhanced error handling
            error_code = "SHAREPOINT_ERROR"
            suggestions = []
            
            if "not found" in result['message'].lower():
                error_code = "FILE_NOT_FOUND"
                suggestions = [
                    "Check if the file exists at the specified path",
                    "Verify the file path is correct",
                    "Ensure the file hasn't been moved or deleted"
                ]
            elif "permission" in result['message'].lower():
                error_code = "PERMISSION_DENIED"
                suggestions = [
                    "Check if you have read permissions for this file",
                    "Contact the file owner or administrator",
                    "Verify your SharePoint access rights"
                ]
            elif "network" in result['message'].lower() or "connection" in result['message'].lower():
                error_code = "NETWORK_ERROR"
                suggestions = [
                    "Check your internet connection",
                    "Verify SharePoint service is available",
                    "Try again in a few minutes"
                ]
            
            return json.dumps({
                "status": "error",
                "error": result['message'],
                "error_code": error_code,
                "path": clean_path,
                "suggestions": suggestions,
                "timestamp": datetime.datetime.now().isoformat()
            })
    except Exception as e:
        logger.error(f"SharePoint read file error: {str(e)}")
        return json.dumps({
            "status": "error",
            "error": f"Unexpected error: {str(e)}",
            "error_code": "UNEXPECTED_ERROR",
            "path": clean_path if 'clean_path' in locals() else path,
            "suggestions": [
                "Check the file path and permissions",
                "Verify SharePoint connection",
                "Contact system administrator if the problem persists"
            ],
            "timestamp": datetime.datetime.now().isoformat()
        })


@mcp.tool()
def sharepoint_list_files(path: str = "", file_type: str = None, sort_by: str = "name", sort_order: str = "asc", max_items: int = 100) -> str:
    """
    List files and directories in SharePoint folder with enhanced filtering and sorting.
    
    Args:
        path: SharePoint folder path (empty for root, e.g., "RFP", "RFP/2024")
        file_type: Filter by file type (e.g., "pdf", "docx", "folder") - default: None (all types)
        sort_by: Sort field ("name", "date", "size", "type") - default: "name"
        sort_order: Sort order ("asc" or "desc") - default: "asc"
        max_items: Maximum number of items to return (1-1000) - default: 100
    
    Returns:
        JSON string with enhanced directory listing including metadata
        
    Examples:
        Basic usage:
        sharepoint_list_files()
        
        List specific folder:
        sharepoint_list_files("RFP")
        
        Filter by file type:
        sharepoint_list_files("RFP", file_type="pdf")
        
        Sort by date (newest first):
        sharepoint_list_files("RFP", sort_by="date", sort_order="desc")
        
        Get more items:
        sharepoint_list_files("RFP", max_items=500)
    """
    # Input validation
    if path and not isinstance(path, str):
        return json.dumps({
            "status": "error",
            "error": "Path must be a string",
            "error_code": "INVALID_PATH_TYPE",
            "suggestions": ["Provide a valid string path"]
        })
    
    # Sanitize and validate path
    clean_path = path.strip() if path else ""
    if '..' in clean_path or clean_path.startswith('/'):
        return json.dumps({
            "status": "error",
            "error": "Invalid path - path traversal not allowed",
            "error_code": "PATH_TRAVERSAL",
            "path": clean_path,
            "suggestions": ["Use relative paths without '..' or absolute paths"]
        })
    
    # Validate parameters
    valid_sort_fields = ["name", "date", "size", "type"]
    if sort_by.lower() not in valid_sort_fields:
        return json.dumps({
            "status": "error",
            "error": f"Invalid sort_by - must be one of: {', '.join(valid_sort_fields)}",
            "error_code": "INVALID_SORT_FIELD",
            "sort_by": sort_by
        })
    
    valid_sort_orders = ["asc", "desc"]
    if sort_order.lower() not in valid_sort_orders:
        return json.dumps({
            "status": "error",
            "error": f"Invalid sort_order - must be one of: {', '.join(valid_sort_orders)}",
            "sort_order": sort_order
        })
    
    if max_items < 1 or max_items > 1000:
        return json.dumps({
            "status": "error",
            "error": "max_items must be between 1 and 1000",
            "error_code": "INVALID_MAX_ITEMS",
            "max_items": max_items
        })
    
    try:
        sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
        from sharepoint_api import SharePointManager  # type: ignore
        
        sharepoint_manager = SharePointManager()
        
        # Get file listing
        logger.info(f"Listing files from SharePoint: {clean_path or 'root'}")
        start_time = datetime.datetime.now()
        
        result = sharepoint_manager.list_files(clean_path)
        
        if result['status'] == 'success':
            files = result['files']
            
            # Process and enhance file information
            enhanced_files = []
            for file_info in files:
                try:
                    # Extract file information
                    file_name = file_info.get('name', 'Unknown')
                    file_path = file_info.get('path', '')
                    is_folder = file_info.get('is_folder', False)
                    
                    # Get file extension and type
                    if not is_folder and '.' in file_name:
                        file_extension = file_name.lower().split('.')[-1]
                        file_type_info = get_file_type_info(file_name)
                    else:
                        file_extension = ''
                        file_type_info = {
                            'extension': '',
                            'is_text_file': False,
                            'is_document': False,
                            'is_image': False,
                            'is_archive': False,
                            'mime_type': 'application/x-directory' if is_folder else 'application/octet-stream',
                            'recommended_max_size_mb': 0
                        }
                    
                    # Build enhanced file info
                    enhanced_file = {
                        'name': file_name,
                        'path': file_path,
                        'is_folder': is_folder,
                        'extension': file_extension,
                        'type': file_type_info['extension'],
                        'mime_type': file_type_info['mime_type'],
                        'is_text_file': file_type_info['is_text_file'],
                        'is_document': file_type_info['is_document'],
                        'is_image': file_type_info['is_image'],
                        'is_archive': file_type_info['is_archive'],
                        'size': file_info.get('size', 0),
                        'size_mb': round(file_info.get('size', 0) / (1024*1024), 2) if file_info.get('size') else 0,
                        'modified_date': file_info.get('modified_date', ''),
                        'created_date': file_info.get('created_date', ''),
                        'owner': file_info.get('owner', ''),
                        'permissions': file_info.get('permissions', '')
                    }
                    
                    # Apply file type filter
                    if file_type:
                        if file_type.lower() == 'folder' and not is_folder:
                            continue
                        elif file_type.lower() != 'folder' and is_folder:
                            continue
                        elif file_type.lower() not in ['folder'] and file_extension.lower() != file_type.lower():
                            continue
                    
                    enhanced_files.append(enhanced_file)
                    
                except Exception as file_error:
                    logger.warning(f"Error processing file {file_info.get('name', 'Unknown')}: {file_error}")
                    continue
            
            # Apply sorting
            reverse_sort = sort_order.lower() == 'desc'
            if sort_by.lower() == 'name':
                enhanced_files.sort(key=lambda x: x['name'].lower(), reverse=reverse_sort)
            elif sort_by.lower() == 'date':
                enhanced_files.sort(key=lambda x: x['modified_date'] or x['created_date'] or '', reverse=reverse_sort)
            elif sort_by.lower() == 'size':
                enhanced_files.sort(key=lambda x: x['size'] or 0, reverse=reverse_sort)
            elif sort_by.lower() == 'type':
                enhanced_files.sort(key=lambda x: x['type'].lower(), reverse=reverse_sort)
            
            # Apply pagination (max_items)
            total_items = len(enhanced_files)
            enhanced_files = enhanced_files[:max_items]
            
            # Calculate processing time
            end_time = datetime.datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            # Build response
            response = {
                "status": "success",
                "path": clean_path or "root",
                "total_items": total_items,
                "returned_items": len(enhanced_files),
                "max_items": max_items,
                "sort_by": sort_by,
                "sort_order": sort_order,
                "file_type_filter": file_type,
                "processing_time_seconds": round(processing_time, 3),
                "items": enhanced_files,
                "timestamp": datetime.datetime.now().isoformat()
            }
            
            # Add summary statistics
            if enhanced_files:
                folders = [f for f in enhanced_files if f['is_folder']]
                documents = [f for f in enhanced_files if f['is_document']]
                images = [f for f in enhanced_files if f['is_image']]
                archives = [f for f in enhanced_files if f['is_archive']]
                
                response["summary"] = {
                    "folders": len(folders),
                    "documents": len(documents),
                    "images": len(images),
                    "archives": len(archives),
                    "other_files": len(enhanced_files) - len(folders) - len(documents) - len(images) - len(archives)
                }
            
            return json.dumps(response)
        else:
            # Enhanced error handling
            error_code = "SHAREPOINT_ERROR"
            suggestions = []
            
            if "not found" in result['message'].lower():
                error_code = "FOLDER_NOT_FOUND"
                suggestions = [
                    "Check if the folder exists at the specified path",
                    "Verify the folder path is correct",
                    "Ensure you have access permissions to this folder"
                ]
            elif "permission" in result['message'].lower():
                error_code = "PERMISSION_DENIED"
                suggestions = [
                    "Check if you have read permissions for this folder",
                    "Contact the folder owner or administrator",
                    "Verify your SharePoint access rights"
                ]
            elif "network" in result['message'].lower() or "connection" in result['message'].lower():
                error_code = "NETWORK_ERROR"
                suggestions = [
                    "Check your internet connection",
                    "Verify SharePoint service is available",
                    "Try again in a few minutes"
                ]
            
            return json.dumps({
                "status": "error",
                "error": result['message'],
                "error_code": error_code,
                "path": clean_path or "root",
                "suggestions": suggestions,
                "timestamp": datetime.datetime.now().isoformat()
            })
    except Exception as e:
        logger.error(f"SharePoint list files error: {str(e)}")
        return json.dumps({
            "status": "error",
            "error": f"Unexpected error: {str(e)}",
            "error_code": "UNEXPECTED_ERROR",
            "path": clean_path if 'clean_path' in locals() else path,
            "suggestions": [
                "Check the folder path and permissions",
                "Verify SharePoint connection",
                "Contact system administrator if the problem persists"
            ],
            "timestamp": datetime.datetime.now().isoformat()
        })


@mcp.tool()
def sharepoint_search_files(query: str, path: str = "", search_type: str = "filename", file_type: str = None, max_results: int = 50, include_content: bool = False) -> str:
    """
    Search for files in SharePoint with enhanced search capabilities.
    
    Args:
        query: Search query (filename pattern or content keywords)
        path: SharePoint folder path to search in (empty for all folders)
        search_type: Search type ("filename", "content", "both") - default: "filename"
        file_type: Filter by file type (e.g., "pdf", "docx", "xlsx") - default: None (all types)
        max_results: Maximum number of results to return (1-200) - default: 50
        include_content: Whether to include file content preview for content searches - default: False
    
    Returns:
        JSON string with enhanced search results including metadata and relevance scores
        
    Examples:
        Basic filename search:
        sharepoint_search_files("RFP 2024")
        
        Content search in specific folder:
        sharepoint_search_files("cloud infrastructure", "RFP/2024", search_type="content")
        
        Search both filename and content:
        sharepoint_search_files("proposal template", search_type="both", file_type="pdf")
        
        Get content previews:
        sharepoint_search_files("technical requirements", include_content=True, max_results=10)
    """
    # Input validation
    if not query or not query.strip():
        return json.dumps({
            "status": "error",
            "error": "Search query cannot be empty",
            "error_code": "EMPTY_QUERY",
            "suggestions": ["Provide a search term or filename pattern"]
        })
    
    # Validate search type
    valid_search_types = ["filename", "content", "both"]
    if search_type.lower() not in valid_search_types:
        return json.dumps({
            "status": "error",
            "error": f"Invalid search_type - must be one of: {', '.join(valid_search_types)}",
            "error_code": "INVALID_SEARCH_TYPE",
            "search_type": search_type
        })
    
    # Validate max_results
    if max_results < 1 or max_results > 200:
        return json.dumps({
            "status": "error",
            "error": "max_results must be between 1 and 200",
            "error_code": "INVALID_MAX_RESULTS",
            "max_results": max_results
        })
    
    # Sanitize and validate path
    clean_path = path.strip() if path else ""
    if '..' in clean_path or clean_path.startswith('/'):
        return json.dumps({
            "status": "error",
            "error": "Invalid path - path traversal not allowed",
            "error_code": "PATH_TRAVERSAL",
            "path": clean_path,
            "suggestions": ["Use relative paths without '..' or absolute paths"]
        })
    
    try:
        sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
        from sharepoint_api import SharePointManager  # type: ignore
        
        sharepoint_manager = SharePointManager()
        
        # Log search parameters
        logger.info(f"SharePoint search: query='{query}', path='{clean_path}', type='{search_type}', file_type='{file_type}', max_results={max_results}")
        start_time = datetime.datetime.now()
        
        # Perform search based on type
        if search_type.lower() == "filename":
            result = sharepoint_manager.search_files(query, clean_path)
        elif search_type.lower() == "content":
            result = sharepoint_manager.search_content(query, clean_path)
        else:  # "both"
            # Combine filename and content search results
            filename_result = sharepoint_manager.search_files(query, clean_path)
            content_result = sharepoint_manager.search_content(query, clean_path)
            
            # Merge results (this would need to be implemented in SharePointManager)
            result = _merge_search_results(filename_result, content_result, query)
        
        if result['status'] == 'success':
            files = result['files']
            
            # Process and enhance search results
            enhanced_results = []
            for file_info in files:
                try:
                    # Extract basic file information
                    file_name = file_info.get('name', 'Unknown')
                    file_path = file_info.get('path', '')
                    is_folder = file_info.get('is_folder', False)
                    
                    # Get file type information
                    if not is_folder and '.' in file_name:
                        file_extension = file_name.lower().split('.')[-1]
                        file_type_info = get_file_type_info(file_name)
                    else:
                        file_extension = ''
                        file_type_info = {
                            'extension': '',
                            'is_text_file': False,
                            'is_document': False,
                            'is_image': False,
                            'is_archive': False,
                            'mime_type': 'application/x-directory' if is_folder else 'application/octet-stream',
                            'recommended_max_size_mb': 0
                        }
                    
                    # Apply file type filter
                    if file_type and file_type.lower() != file_extension.lower():
                        continue
                    
                    # Calculate relevance score based on search type
                    relevance_score = _calculate_relevance_score(file_info, query, search_type)
                    
                    # Build enhanced file info
                    enhanced_file = {
                        'name': file_name,
                        'path': file_path,
                        'is_folder': is_folder,
                        'extension': file_extension,
                        'type': file_type_info['extension'],
                        'mime_type': file_type_info['mime_type'],
                        'is_text_file': file_type_info['is_text_file'],
                        'is_document': file_type_info['is_document'],
                        'is_image': file_type_info['is_image'],
                        'is_archive': file_type_info['is_archive'],
                        'size': file_info.get('size', 0),
                        'size_mb': round(file_info.get('size', 0) / (1024*1024), 2) if file_info.get('size') else 0,
                        'modified_date': file_info.get('modified_date', ''),
                        'created_date': file_info.get('created_date', ''),
                        'owner': file_info.get('owner', ''),
                        'relevance_score': relevance_score,
                        'search_matches': file_info.get('search_matches', [])
                    }
                    
                    # Add content preview for content searches
                    if include_content and search_type.lower() in ["content", "both"] and file_type_info['is_text_file']:
                        try:
                            content_preview = _get_content_preview(file_path, query)
                            if content_preview:
                                enhanced_file['content_preview'] = content_preview
                        except Exception as preview_error:
                            logger.warning(f"Could not get content preview for {file_path}: {preview_error}")
                    
                    enhanced_results.append(enhanced_file)
                    
                except Exception as file_error:
                    logger.warning(f"Error processing search result {file_info.get('name', 'Unknown')}: {file_error}")
                    continue
            
            # Sort by relevance score (highest first)
            enhanced_results.sort(key=lambda x: x['relevance_score'], reverse=True)
            
            # Apply max_results limit
            total_results = len(enhanced_results)
            enhanced_results = enhanced_results[:max_results]
            
            # Calculate processing time
            end_time = datetime.datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            
            # Build response
            response = {
                "status": "success",
                "query": query.strip(),
                "path": clean_path or "all folders",
                "search_type": search_type,
                "file_type_filter": file_type,
                "total_results": total_results,
                "returned_results": len(enhanced_results),
                "max_results": max_results,
                "processing_time_seconds": round(processing_time, 3),
                "results": enhanced_results,
                "timestamp": datetime.datetime.now().isoformat()
            }
            
            # Add search suggestions if no results
            if total_results == 0:
                response["suggestions"] = [
                    "Try different keywords or search terms",
                    "Check if the file type filter is too restrictive",
                    "Search in a different folder path",
                    "Try a broader search query"
                ]
            
            return json.dumps(response)
        else:
            # Enhanced error handling
            error_code = "SHAREPOINT_ERROR"
            suggestions = []
            
            if "not found" in result['message'].lower():
                error_code = "NO_RESULTS"
                suggestions = [
                    "Try different search terms",
                    "Check the folder path",
                    "Verify file permissions"
                ]
            elif "permission" in result['message'].lower():
                error_code = "PERMISSION_DENIED"
                suggestions = [
                    "Check your SharePoint access permissions",
                    "Contact the folder owner or administrator"
                ]
            elif "network" in result['message'].lower() or "connection" in result['message'].lower():
                error_code = "NETWORK_ERROR"
                suggestions = [
                    "Check your internet connection",
                    "Verify SharePoint service is available"
                ]
            
            return json.dumps({
                "status": "error",
                "error": result['message'],
                "error_code": error_code,
                "query": query.strip(),
                "path": clean_path or "all folders",
                "suggestions": suggestions,
                "timestamp": datetime.datetime.now().isoformat()
            })
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})

@mcp.tool()
def extract_pdf_text(path: str, pages: str = "all", clean_text: bool = True, preserve_structure: bool = True, extract_tables: bool = False, max_pages: int = 50) -> str:
    """
    Extract and process text content from PDF files in SharePoint with enhanced features for RFP analysis.
    
    Args:
        path: SharePoint path to PDF file
        pages: Page range to extract ("all", "1-5", "1,3,5", "first", "last") - default: "all"
        clean_text: Whether to clean and format extracted text - default: True
        preserve_structure: Whether to preserve document structure (headings, sections) - default: True
        extract_tables: Whether to attempt table extraction - default: False
        max_pages: Maximum number of pages to process (safety limit) - default: 50
    
    Returns:
        JSON string with enhanced extracted text and metadata
        
    Examples:
        Basic extraction:
        extract_pdf_text("RFP/2024/proposal.pdf")
        
        Extract specific pages:
        extract_pdf_text("RFP/2024/proposal.pdf", pages="1-5")
        
        Extract with table detection:
        extract_pdf_text("RFP/2024/proposal.pdf", extract_tables=True)
        
        Extract first 10 pages only:
        extract_pdf_text("RFP/2024/proposal.pdf", pages="first", max_pages=10)
        
        Raw extraction (no cleaning):
        extract_pdf_text("RFP/2024/proposal.pdf", clean_text=False, preserve_structure=False)
    """
    # Input validation
    if not path or not path.strip():
        return json.dumps({
            "status": "error",
            "error": "PDF file path cannot be empty",
            "error_code": "EMPTY_PATH",
            "suggestions": ["Provide a valid SharePoint path to the PDF file"]
        })
    
    # Validate max_pages
    if max_pages < 1 or max_pages > 500:
        return json.dumps({
            "status": "error",
            "error": "max_pages must be between 1 and 500",
            "error_code": "INVALID_MAX_PAGES",
            "max_pages": max_pages
        })
    
    # Sanitize and validate path
    clean_path = path.strip().replace('\\', '/')
    if '..' in clean_path or clean_path.startswith('/'):
        return json.dumps({
            "status": "error",
            "error": "Invalid file path - path traversal not allowed",
            "error_code": "PATH_TRAVERSAL",
            "path": clean_path,
            "suggestions": ["Use relative paths without '..' or absolute paths"]
        })
    
    # Validate file extension
    if not clean_path.lower().endswith('.pdf'):
        return json.dumps({
            "status": "error",
            "error": "File must be a PDF",
            "error_code": "INVALID_FILE_TYPE",
            "path": clean_path,
            "suggestions": ["Provide a path to a PDF file"]
        })
    
    try:
        sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))
        from sharepoint_api import SharePointManager  # type: ignore
        import PyPDF2
        from io import BytesIO
        
        sharepoint_manager = SharePointManager()
        
        # Log extraction parameters
        logger.info(f"PDF text extraction: path='{clean_path}', pages='{pages}', clean_text={clean_text}, preserve_structure={preserve_structure}, extract_tables={extract_tables}, max_pages={max_pages}")
        start_time = datetime.datetime.now()
        
        file_result = sharepoint_manager.get_file_content(clean_path, binary=True)
        
        if file_result['status'] != 'success':
            return json.dumps({
                "status": "error",
                "error": file_result['message'],
                "error_code": "SHAREPOINT_ERROR",
                "path": clean_path
            })
        
        # Process PDF
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_result['content']))
        total_pages = len(pdf_reader.pages)
        
        # Determine pages to extract
        pages_to_extract = _parse_page_range(pages, total_pages, max_pages)
        
        if not pages_to_extract:
            return json.dumps({
                "status": "error",
                "error": f"No valid pages to extract. Total pages: {total_pages}, requested: {pages}",
                "error_code": "INVALID_PAGE_RANGE",
                "total_pages": total_pages,
                "requested_pages": pages
            })
        
        # Extract text from specified pages
        extracted_text = ""
        page_texts = []
        tables_found = []
        
        for page_num in pages_to_extract:
            try:
                page = pdf_reader.pages[page_num - 1]  # Convert to 0-based index
                page_text = page.extract_text()
                
                if clean_text:
                    page_text = _clean_pdf_text(page_text)
                
                if preserve_structure:
                    page_text = _preserve_document_structure(page_text)
                
                if extract_tables:
                    table_data = _extract_table_data(page)
                    if table_data:
                        tables_found.append({
                            "page": page_num,
                            "tables": table_data
                        })
                
                page_texts.append({
                    "page": page_num,
                    "text": page_text,
                    "char_count": len(page_text)
                })
                
                extracted_text += page_text + "\n\n"
                
            except Exception as page_error:
                logger.warning(f"Error extracting page {page_num}: {page_error}")
                page_texts.append({
                    "page": page_num,
                    "text": f"[Error extracting page {page_num}]",
                    "char_count": 0,
                    "error": str(page_error)
                })
        
        # Calculate processing time
        end_time = datetime.datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        
        # Build response
        response = {
            "status": "success",
            "path": clean_path,
            "total_pages": total_pages,
            "extracted_pages": len(pages_to_extract),
            "pages_requested": pages,
            "pages_extracted": pages_to_extract,
            "text": extracted_text.strip(),
            "total_characters": len(extracted_text.strip()),
            "clean_text": clean_text,
            "preserve_structure": preserve_structure,
            "extract_tables": extract_tables,
            "processing_time_seconds": round(processing_time, 3),
            "page_details": page_texts,
            "timestamp": datetime.datetime.now().isoformat()
        }
        
        if tables_found:
            response["tables_found"] = tables_found
            response["table_count"] = sum(len(t["tables"]) for t in tables_found)
        
        # Add extraction statistics
        response["statistics"] = {
            "avg_chars_per_page": round(len(extracted_text.strip()) / len(pages_to_extract), 2) if pages_to_extract else 0,
            "pages_with_errors": len([p for p in page_texts if "error" in p]),
            "extraction_success_rate": round((len(pages_to_extract) - len([p for p in page_texts if "error" in p])) / len(pages_to_extract) * 100, 1) if pages_to_extract else 0
        }
        
        return json.dumps(response)
    except Exception as e:
        return json.dumps({"status": "error", "error": str(e)})



@mcp.prompt(title="Identifying the Documents")
def Step1_Identifying_documents():
    """Identify and categorize RFP/RFI/RFQ documents from SharePoint folders."""
    return f"""Analyze the selected SharePoint folder:

- Use sharepoint_list_files to list files
- Use sharepoint_read_file with preview_lines=50 for PDFs/text files (only first 50 lines)

Categorize each file into:
1. 📘 **Primary Documents** — RFP/RFI/RFQ content (scope, requirements, evaluation criteria)  
2. 📄 **Supporting Documents** — Appendices, attachments, specs  
3. 📝 **Other Documents** — Forms, templates, misc files  

**Rules:**  
- Read only previews (no full documents)  
- No assumptions about structure/content  
- Output filenames + one-line explanation only  
- End with "ANALYSIS COMPLETE"  

**Format:**  

📘 **Primary Documents**  
- filename1.pdf - Contains RFP requirements  

📄 **Supporting Documents**  
- filename2.pdf - Technical appendix  

📝 **Other Documents**  
- filename3.pdf - Response form  

ANALYSIS COMPLETE
"""


@mcp.prompt(title="Step-2: Executive Summary of Procurement Document")
def Step2_summarize_documents():
    """Concise executive summary of RFP/RFQ/RFI documents."""
    return f"""
For each document, provide:

### 📄 Document: [Document Name]

#### 🔹 Overview
3–5 sentence plain-English summary: who issued it, what’s needed, why, and expected response.

#### 🧩 Key Details
- **Deadline:** [Date + Time]  
- **Project Dates:** [Start/End]  
- **Budget/Value:** [If stated]  
- **Response Format:** [PDF, portal, etc.]  
- **Location:** [Region/Remote]  
- **Eligibility:** [Certs, licenses, limits]  
- **Scope:** [Main tasks/deliverables]  

#### 📊 Evaluation
How responses will be scored/selected (include weights if given).

#### 📝 Notes
- **Risks:** Red flags or challenges  
- **Opportunities:** Competitive advantages / differentiators  

#### 📞 Contact
- **Contact:** [Name, role, email/phone]  
- **Submission:** [Portal, email, etc.]  

### 🤔 Next Step
Ask: *Would you like a strategic assessment or Go/No-Go recommendation?*

⚠️ Only include what’s explicitly stated — no assumptions.
"""



@mcp.prompt(title="Step-3 : Go/No-Go Recommendation")
def Step3_go_no_go_recommendation() -> str:
    """Generate an exec-style Go/No-Go matrix using BroadAxis knowledge base."""
    return """
You are BroadAxis-AI. You have already generated executive summaries (Step-2). Now evaluate BroadAxis’s fit using the knowledge base and present findings in a concise matrix.

---

### 📊 Capability Match Matrix

| Requirement              | BroadAxis Capability / Evidence        | Match Level (✅/⚠️/❌) | Notes / Gaps |
|---------------------------|-----------------------------------------|------------------------|--------------|
| [Requirement 1]           | [Relevant project/case study/skill]     | ✅                     | -            |
| [Requirement 2]           | [Partial experience / limited coverage] | ⚠️                     | Needs partner support |
| [Requirement 3]           | [No evidence found]                     | ❌                     | Missing certs |

---

### 📋 GO/NO-GO RECOMMENDATION
- **Recommendation:** [GO / NO-GO / NEED INFO]  
- **Confidence:** [HIGH / MEDIUM / LOW]  
- **Top Factors:** [3–5 key drivers]  
- **Next Steps (if GO):** [Actions required]  
- **Missing Info (if NEED INFO):** [What’s needed for decision]  

---

⚠️ Use only verified knowledge base results. Do not guess. Flag missing data clearly.
"""

@mcp.prompt(title="Step-4 : Generate Proposal or Capability Statement")
def Step4_generate_capability_statement() -> str:
    return """
You are BroadAxis-AI, an assistant trained to generate high-quality capability statements and proposal documents for RFP and RFQ responses.
The user has either uploaded an opportunity document or requested a formal proposal. Use all available information from:

- SharePoint documents (RFP/RFQ)
- Internal knowledge base
- Prior summaries or analyses already provided

---

### 🧠 Instructions

- Do not invent names, projects, or facts.
- Use internal knowledge to populate all relevant content.
- Leave placeholders where personal or business info is not available.
- Maintain professional, confident, and compliant tone.

If this proposal is meant to be saved, offer to generate a PDF or Word version using the appropriate tool.

"""

@mcp.prompt(title="Step-5 : Fill in Missing Information")
def Step5_fill_missing_information() -> str:
    return """
You are BroadAxis-AI, an intelligent assistant designed to fill in missing fields, answer RFP/RFQ questions, and complete response templates **strictly using verified information**.
Your task is to **complete the missing sections** on the fillable documents which you have identified previously with reliable information from:

1. Internal knowledge base
2. The uploaded document itself
3. Prior chat context (if available)
---
### 🧠 RULES (Strict Compliance)

- ❌ **DO NOT invent or hallucinate** company details, financials, certifications, team names, or client info.
- ❌ **DO NOT guess** values you cannot verify.
- 🔐 If the question involves **personal, legal, or confidential information**, **do not fill it**.
- ✅ Use internal knowledge only when it clearly answers the field.
---
### ✅ Final Instruction
Only fill what you can verify using internal knowledge and uploaded content. Leave everything else with clear, professional placeholders.

"""


if __name__ == "__main__":
    import sys
    import logging

    # Set up logging for debugging
    logging.basicConfig(level=logging.INFO)

    try:
        # Run the MCP server synchronously
        mcp.run()
    except KeyboardInterrupt:
        print("Server stopped", file=sys.stderr)
    except Exception as e:
        print(f"Server error: {e}", file=sys.stderr)
        sys.exit(1)